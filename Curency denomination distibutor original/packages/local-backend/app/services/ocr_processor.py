"""
OCR Processing Service for Bulk Upload - Rebuilt from Scratch

Handles text extraction from various file formats:
- CSV files (direct parsing, no OCR needed)
- Images (JPG, PNG, TIFF, BMP) - Tesseract OCR
- PDFs (text extraction + OCR for scanned PDFs) - PyMuPDF + pdf2image + Tesseract
- Word documents (.docx) - python-docx

Fully offline after dependencies are installed.
"""

import os
import re
import io
import tempfile
from pathlib import Path
from typing import List, Dict, Any, Optional
from decimal import Decimal, InvalidOperation
import logging

# Configure logging
logger = logging.getLogger(__name__)

# Import optional OCR dependencies
try:
    import pytesseract
    from PIL import Image
    HAS_TESSERACT = True
except ImportError:
    HAS_TESSERACT = False
    logger.warning("Tesseract OCR not available")

try:
    import fitz  # PyMuPDF
    HAS_PYMUPDF = True
except ImportError:
    HAS_PYMUPDF = False
    logger.warning("PyMuPDF not available")

try:
    from pdf2image import convert_from_bytes
    HAS_PDF2IMAGE = True
except ImportError:
    HAS_PDF2IMAGE = False
    logger.warning("pdf2image not available")

try:
    import docx
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    logger.warning("python-docx not available")


class OCRProcessor:
    """
    Handles OCR and text extraction from multiple file formats.
    Enhanced with intelligent parsing and smart defaults.
    """
    
    def __init__(self, default_currency: str = 'INR', default_mode: str = 'greedy'):
        """Initialize OCR processor with defaults."""
        self.supported_image_formats = {'.jpg', '.jpeg', '.png', '.tiff', '.tif', '.bmp', '.gif', '.webp'}
        self.supported_pdf_formats = {'.pdf'}
        self.supported_word_formats = {'.docx', '.doc'}
        
        # Smart defaults
        self.default_currency = default_currency
        self.default_mode = default_mode
        
        logger.info(f"OCR Processor initialized (default currency: {default_currency}, default mode: {default_mode})")
    
    def check_dependencies(self) -> Dict[str, bool]:
        """Check which OCR dependencies are available."""
        return {
            'tesseract': HAS_TESSERACT,
            'pymupdf': HAS_PYMUPDF,
            'pdf2image': HAS_PDF2IMAGE,
            'docx': HAS_DOCX
        }
    
    def process_file(self, file_data: bytes, filename: str) -> List[Dict[str, Any]]:
        """
        Process uploaded file and extract structured data.
        
        Args:
            file_data: Raw file bytes
            filename: Original filename
        
        Returns:
            List of dicts with keys: row_number, amount, currency, optimization_mode
        
        Raises:
            ValueError: If file format not supported or extraction fails
        """
        file_ext = Path(filename).suffix.lower()
        
        logger.info(f"Processing file: {filename} (size: {len(file_data)} bytes, type: {file_ext})")
        
        # Route to appropriate processor based on file type
        if file_ext in self.supported_image_formats:
            return self._process_image(file_data, filename)
        elif file_ext in self.supported_pdf_formats:
            return self._process_pdf(file_data, filename)
        elif file_ext in self.supported_word_formats:
            return self._process_word(file_data, filename)
        else:
            raise ValueError(f"Unsupported file format: {file_ext}")
    
    def _process_image(self, file_data: bytes, filename: str) -> List[Dict[str, Any]]:
        """Extract text from image using Tesseract OCR."""
        if not HAS_TESSERACT:
            raise ValueError("Tesseract OCR not installed. Run: install_ocr_simple.ps1")
        
        try:
            # Load image from bytes
            image = Image.open(io.BytesIO(file_data))
            
            logger.info(f"Image loaded: {image.size}, mode: {image.mode}")
            
            # Perform OCR
            extracted_text = pytesseract.image_to_string(image)
            
            logger.debug(f"Extracted text ({len(extracted_text)} chars):\n{extracted_text[:500]}")
            
            # Parse extracted text into structured rows
            return self._parse_text_to_rows(extracted_text)
            
        except Exception as e:
            logger.error(f"Image OCR failed: {str(e)}")
            raise ValueError(f"Failed to process image: {str(e)}")
    
    def _process_pdf(self, file_data: bytes, filename: str) -> List[Dict[str, Any]]:
        """Extract text from PDF (text-based or scanned)."""
        if not HAS_PYMUPDF:
            raise ValueError("PyMuPDF not installed. Run: install_ocr_simple.ps1")
        
        try:
            # Try text extraction first (faster for text-based PDFs)
            pdf_document = fitz.open(stream=file_data, filetype="pdf")
            extracted_text = ""
            
            for page_num in range(len(pdf_document)):
                page = pdf_document[page_num]
                extracted_text += page.get_text()
            
            pdf_document.close()
            
            logger.debug(f"PDF text extracted ({len(extracted_text)} chars)")
            
            # If no text extracted, try OCR on scanned PDF
            if len(extracted_text.strip()) < 50:
                logger.info("PDF appears to be scanned, attempting OCR")
                return self._process_scanned_pdf(file_data, filename)
            
            # Parse extracted text
            return self._parse_text_to_rows(extracted_text)
            
        except Exception as e:
            logger.error(f"PDF processing failed: {str(e)}")
            raise ValueError(f"Failed to process PDF: {str(e)}")
    
    def _process_scanned_pdf(self, file_data: bytes, filename: str) -> List[Dict[str, Any]]:
        """Process scanned PDF using OCR."""
        if not HAS_PDF2IMAGE or not HAS_TESSERACT:
            raise ValueError("pdf2image or Tesseract not installed for scanned PDFs")
        
        try:
            # Convert PDF pages to images
            images = convert_from_bytes(file_data)
            logger.info(f"Converted PDF to {len(images)} images")
            
            extracted_text = ""
            for idx, image in enumerate(images):
                logger.debug(f"OCR on page {idx + 1}/{len(images)}")
                page_text = pytesseract.image_to_string(image)
                extracted_text += page_text + "\n"
            
            logger.debug(f"Total extracted text: {len(extracted_text)} chars")
            
            return self._parse_text_to_rows(extracted_text)
            
        except Exception as e:
            logger.error(f"Scanned PDF OCR failed: {str(e)}")
            raise ValueError(f"Failed to OCR scanned PDF: {str(e)}")
    
    def _process_word(self, file_data: bytes, filename: str) -> List[Dict[str, Any]]:
        """Extract text from Word document."""
        if not HAS_DOCX:
            raise ValueError("python-docx not installed. Run: install_ocr_simple.ps1")
        
        try:
            # Load Word document from bytes
            doc = docx.Document(io.BytesIO(file_data))
            
            # Extract all text from paragraphs
            extracted_text = "\n".join([para.text for para in doc.paragraphs])
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text for cell in row.cells])
                    extracted_text += "\n" + row_text
            
            logger.debug(f"Word document text extracted ({len(extracted_text)} chars)")
            
            return self._parse_text_to_rows(extracted_text)
            
        except Exception as e:
            logger.error(f"Word document processing failed: {str(e)}")
            raise ValueError(f"Failed to process Word document: {str(e)}")
    
    def _parse_text_to_rows(self, text: str) -> List[Dict[str, Any]]:
        """
        Parse extracted text into structured rows.
        
        Expected format examples:
        1. CSV-like: "125.50, USD, greedy"
        2. Tabular: "Amount | Currency | Mode"
        3. Natural: "Amount: 125.50 Currency: USD Mode: greedy"
        """
        rows = []
        lines = text.strip().split('\n')
        
        logger.info(f"Parsing {len(lines)} lines of text")
        
        # Detect header row (skip it)
        header_keywords = ['amount', 'currency', 'mode', 'optimization']
        
        for line_num, line in enumerate(lines, start=1):
            line = line.strip()
            
            if not line:
                continue
            
            # Skip header rows
            if any(keyword in line.lower() for keyword in header_keywords):
                if line_num == 1 or '|' in line or '-' * 3 in line:
                    logger.debug(f"Skipping header line {line_num}: {line}")
                    continue
            
            # Try to parse the line
            try:
                parsed_row = self._parse_line(line, line_num)
                if parsed_row:
                    rows.append(parsed_row)
                    logger.debug(f"Line {line_num} parsed: {parsed_row}")
            except Exception as e:
                logger.warning(f"Failed to parse line {line_num}: {line} - {str(e)}")
                continue
        
        logger.info(f"Successfully parsed {len(rows)} rows from {len(lines)} lines")
        
        return rows
    
    def _parse_line(self, line: str, line_number: int) -> Optional[Dict[str, Any]]:
        """
        ENHANCED: Parse a single line with intelligent extraction and smart defaults.
        
        Handles ANY format:
        - CSV: "125.50, USD, greedy" or "125.50, USD" or "125.50"
        - Pipe: "125.50 | USD | greedy"
        - Tabular: "125.50    USD    greedy"
        - Natural: "Amount: 125.50 Currency: USD Mode: greedy"
        - Mixed: "125.50 USD" or "1000 INR greedy" or just "5000"
        
        Smart Defaults:
        - Missing currency → uses system default
        - Missing mode → uses 'greedy'
        """
        # Extract amount first (required)
        amount = self._smart_extract_amount(line)
        if not amount:
            return None  # No valid amount found
        
        # Extract currency (optional, defaults to system default)
        currency = self._smart_extract_currency(line)
        if not currency:
            currency = self.default_currency
            logger.debug(f"Line {line_number}: No currency found, using default: {currency}")
        
        # Extract mode (optional, defaults to greedy)
        mode = self._smart_extract_mode(line)
        if not mode:
            mode = self.default_mode
            logger.debug(f"Line {line_number}: No mode found, using default: {mode}")
        
        return {
            'row_number': line_number,
            'amount': amount,
            'currency': currency,
            'optimization_mode': mode
        }
    
    def _smart_extract_amount(self, text: str) -> str:
        """ENHANCED: Intelligently extract amount from any text format."""
        # Strategy 1: Look for explicit amount labels
        amount_match = re.search(r'(?:amount|amt|value|price|total)[:\s]*([0-9.,E+-]+)', text, re.IGNORECASE)
        if amount_match:
            return self._clean_amount(amount_match.group(1))
        
        # Strategy 2: Find first number in the line (most common)
        number_match = re.search(r'([0-9.,E+-]+)', text)
        if number_match:
            return self._clean_amount(number_match.group(1))
        
        return ''
    
    def _clean_amount(self, text: str) -> str:
        """Clean and normalize amount string."""
        # Remove currency symbols and extra whitespace
        cleaned = re.sub(r'[₹$€£¥,\s]', '', text)
        
        # Handle scientific notation (e.g., 1.23E+10)
        if 'E' in cleaned.upper():
            try:
                float_val = float(cleaned)
                return str(float_val)
            except ValueError:
                pass
        
        return cleaned
    
    def _extract_amount(self, text: str) -> str:
        """Legacy method - redirects to smart extraction."""
        return self._smart_extract_amount(text)
    
    def _smart_extract_currency(self, text: str) -> str:
        """ENHANCED: Intelligently extract currency from any text format."""
        # Strategy 1: Look for currency symbols and names first (most specific)
        text_lower = text.lower()
        if '₹' in text or 'rs.' in text_lower or 'rupee' in text_lower:
            return 'INR'
        if '$' in text or 'dollar' in text_lower:
            return 'USD'
        if '€' in text or 'euro' in text_lower:
            return 'EUR'
        if '£' in text or 'pound' in text_lower or 'sterling' in text_lower:
            return 'GBP'
        
        # Strategy 2: Look for explicit currency labels
        currency_label_match = re.search(r'(?:currency|cur)[:\s]*([A-Z]{3}|\w+)', text, re.IGNORECASE)
        if currency_label_match:
            return self._normalize_currency(currency_label_match.group(1))
        
        # Strategy 3: Look for 3-letter currency codes anywhere
        currency_code_match = re.search(r'\b([A-Z]{3})\b', text.upper())
        if currency_code_match:
            code = currency_code_match.group(1)
            # Filter out common non-currency 3-letter words
            if code not in ['THE', 'AND', 'FOR', 'ARE', 'YOU', 'NOT', 'BUT', 'CAN', 'ALL']:
                return self._normalize_currency(code)
        
        return ''  # No currency found, will use default
    
    def _normalize_currency(self, text: str) -> str:
        """Normalize currency names and codes."""
        # Common currency name corrections
        corrections = {
            'RUPEE': 'INR', 'RUPEES': 'INR', 'RS': 'INR', 'INDIAN': 'INR',
            'DOLLAR': 'USD', 'DOLLARS': 'USD', 'BUCK': 'USD', 'BUCKS': 'USD',
            'EURO': 'EUR', 'EUROS': 'EUR',
            'POUND': 'GBP', 'POUNDS': 'GBP', 'STERLING': 'GBP',
            'YEN': 'JPY', 'JAPANESE': 'JPY',
            'YUAN': 'CNY', 'RENMINBI': 'CNY',
            'CANADIAN': 'CAD', 'LOONIE': 'CAD'
        }
        
        text_upper = text.upper().strip()
        return corrections.get(text_upper, text_upper if len(text_upper) == 3 else '')
    
    def _extract_currency(self, text: str) -> str:
        """Legacy method - redirects to smart extraction."""
        return self._smart_extract_currency(text)
    
    def _smart_extract_mode(self, text: str) -> str:
        """ENHANCED: Intelligently extract optimization mode from any text format."""
        if not text:
            return ''  # Will use default
        
        text_lower = text.lower().strip()
        
        # Strategy 1: Look for explicit mode labels
        mode_label_match = re.search(r'(?:mode|method|optimization|opt|strategy)[:\s]*(\w+)', text_lower)
        if mode_label_match:
            mode_text = mode_label_match.group(1)
            return self._normalize_mode(mode_text)
        
        # Strategy 2: Look for mode keywords anywhere in text
        return self._normalize_mode(text_lower)
    
    def _normalize_mode(self, text: str) -> str:
        """Normalize mode text to valid optimization mode."""
        if not text:
            return ''
        
        text_lower = text.lower().strip()
        
        # Valid modes
        valid_modes = ['greedy', 'balanced', 'minimize_large', 'minimize_small']
        
        # Direct match
        if text_lower in valid_modes:
            return text_lower
        
        # Partial matches and aliases
        if 'bal' in text_lower or 'even' in text_lower or 'equal' in text_lower:
            return 'balanced'
        if 'large' in text_lower or 'big' in text_lower or 'max' in text_lower:
            return 'minimize_large'
        if 'small' in text_lower or 'little' in text_lower or 'tiny' in text_lower:
            return 'minimize_small'
        if 'greed' in text_lower or 'fast' in text_lower or 'quick' in text_lower:
            return 'greedy'
        
        return ''  # No mode found, will use default
    
    def _extract_mode(self, text: str) -> str:
        """Legacy method - redirects to smart extraction."""
        result = self._smart_extract_mode(text)
        return result if result else self.default_mode
    
    def _looks_like_amount(self, text: str) -> bool:
        """Check if text looks like a numeric amount."""
        # Remove common separators
        cleaned = text.replace(',', '').replace(' ', '').replace('₹', '').replace('$', '')
        
        # Check if it's a number (including scientific notation)
        try:
            float(cleaned)
            return True
        except ValueError:
            return False


# Singleton instance
_ocr_processor_instance = None

def get_ocr_processor() -> OCRProcessor:
    """Get singleton OCR processor instance."""
    global _ocr_processor_instance
    if _ocr_processor_instance is None:
        _ocr_processor_instance = OCRProcessor()
    return _ocr_processor_instance
